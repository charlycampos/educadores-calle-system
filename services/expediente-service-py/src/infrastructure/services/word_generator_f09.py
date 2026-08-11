"""Genera el Word del Informe Situacional.

El educador lo descarga, lo revisa, lo firma y lo tramita por el SGD:
"les va a generar este Word que está acá... lo llena y se les genera nomás el
Word y ustedes ya lo tramitan por SGD" (reunión 05/08/2026).

La estructura sigue al modelo oficial (informe de los hermanos Ruiz Culqui):
membrete, número correlativo, ocho secciones en números romanos y firma.

`python-docx` se importa dentro de la función a propósito. Si la librería no
está instalada, el que falla es este endpoint y no el arranque del servicio
entero.
"""

from typing import Optional

MESES_FASE = {1: 3, 2: 15, 3: 6}

_MEMBRETE = "“Año de la recuperación y consolidación de la economía peruana”"
_PIE_INSTITUCIONAL = [
    "Av. San Martín 685, Pueblo Libre",
    "www.inabif.gob.pe   Lima 21, Perú",
    "T. 417-6720",
]


def _viñetas(doc, texto: Optional[str]):
    """Una línea del textarea = una viñeta. Las líneas en blanco se ignoran."""
    for linea in (texto or "").splitlines():
        limpia = linea.strip().lstrip("•-–").strip()
        if limpia:
            doc.add_paragraph(limpia, style="List Bullet")


def _parrafos(doc, texto: Optional[str]):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for bloque in (texto or "").split("\n\n"):
        limpio = bloque.strip()
        if limpio:
            p = doc.add_paragraph(limpio)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _titulo_seccion(doc, texto: str):
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(11)


def generate_f09_docx(informe: dict, nnas: list, educador: dict, salida: str) -> str:
    """Escribe el .docx en `salida` y devuelve la ruta.

    `nnas` es la lista de NNA que cubre el informe: son varios cuando son
    hermanos, y cada uno lleva su propio bloque en la sección I.
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    estilo = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(11)
    for seccion in doc.sections:
        seccion.top_margin = Cm(2.5)
        seccion.bottom_margin = Cm(2.5)
        seccion.left_margin = Cm(3)
        seccion.right_margin = Cm(2.5)

    # ── Membrete ──
    enc = doc.add_paragraph()
    enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = enc.add_run(_MEMBRETE)
    run.italic = True
    run.font.size = Pt(9)

    for linea in _PIE_INSTITUCIONAL:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(linea)
        r.font.size = Pt(8)

    # ── Número del informe ──
    tit = doc.add_paragraph()
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tit.add_run(f"INFORME SITUACIONAL {informe.get('codigo_informe') or ''}".strip())
    r.bold = True
    r.font.size = Pt(12)

    # ── I. Datos generales ──
    _titulo_seccion(doc, "I. DATOS GENERALES DE LA NIÑA, NIÑO O ADOLESCENTE")
    for nna in nnas:
        for etiqueta, clave in [
            ("Nombres Y Apellidos", "nombre_completo"),
            ("Edad", "edad"),
            ("Lugar De Nacimiento", "lugar_nacimiento"),
            ("Fecha De Nacimiento", "fecha_nacimiento"),
            ("Documento De Identificación", "numero_doc"),
            ("Grado De Instrucción", "grado_instruccion"),
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run(f"{etiqueta} : ").bold = True
            p.add_run(str(nna.get(clave) or "---"))
        doc.add_paragraph()

    for etiqueta, clave in [
        ("Perfil del NNA", "perfil"),
        ("Dirección", "direccion"),
        ("Referencia", "referencia"),
        ("Referente familiar de contacto", "referente"),
        ("Teléfono", "telefono"),
        ("Fecha de informe", "fecha_informe"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"{etiqueta} : ").bold = True
        p.add_run(str(informe.get(clave) or "---"))

    doc.add_paragraph()

    # ── II a VIII ──
    _titulo_seccion(doc, "II. ANTECEDENTES DEL CASO:")
    _viñetas(doc, informe.get("antecedentes"))

    _titulo_seccion(doc, "III. ACCIONES REALIZADAS:")
    _viñetas(doc, informe.get("estrategias"))

    _titulo_seccion(doc, "IV. SITUACIÓN FAMILIAR:")
    _parrafos(doc, informe.get("situacion_familiar"))

    _titulo_seccion(doc, "V. INDICADORES DE VULNERABILIDAD")
    _viñetas(doc, informe.get("indicadores_vulnerab"))

    _titulo_seccion(doc, "VI. PLAN DE INTERVENCIÓN INDIVIDUAL")
    for fase in (1, 2, 3):
        contenido = informe.get(f"pii_fase{fase}")
        if not (contenido or "").strip():
            continue
        p = doc.add_paragraph()
        r = p.add_run(f"Fase {fase} ({MESES_FASE[fase]} meses)")
        r.bold = True
        _viñetas(doc, contenido)

    _titulo_seccion(doc, "VII. APRECIACIÓN PROFESIONAL")
    _parrafos(doc, informe.get("conclusiones"))

    _titulo_seccion(doc, "VIII. RECOMENDACIÓN")
    _parrafos(doc, informe.get("recomendaciones"))

    # ── Cierre y firma ──
    doc.add_paragraph()
    for linea in ("Es todo cuanto tengo que informar", "Atentamente."):
        p = doc.add_paragraph(linea)
        p.paragraph_format.left_indent = Cm(6)

    doc.add_paragraph()
    doc.add_paragraph()

    firma = doc.add_paragraph("..………………………………………………………")
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for texto in [
        educador.get("nombre") or "",
        educador.get("cargo") or "Educador/a de calle",
        f"Correo: {educador.get('correo')}" if educador.get("correo") else "",
    ]:
        if not texto:
            continue
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)

    doc.save(salida)
    return salida


# ── Formato de los datos que salen de la base ────────────────────────────────
# Mismo catálogo que usa el cliente (client/src/data/catalogos-sec.ts), tomado
# del Diccionario_de_datos_SEC_2026.

_GRADO = {
    "1": "Inicial", "2": "1ro primaria", "3": "2do primaria", "4": "3ro primaria",
    "5": "4to primaria", "6": "5to primaria", "7": "6to primaria",
    "8": "1ro secundaria", "9": "2do secundaria", "10": "3ro secundaria",
    "11": "4to secundaria", "12": "5to secundaria", "99": "",
}

_NIVEL = {
    "1": "Sin nivel", "2": "Inicial",
    "3": "Primaria Incompleta", "4": "Primaria Completa",
    "5": "Secundaria Incompleta", "6": "Secundaria Completa",
    "7": "Superior No Univ. Incompleta", "8": "Superior No Univ. Completa",
    "9": "Superior Univ. Incompleto", "10": "Superior Univ. Completo",
    "11": "Básica Especial",
}

_MESES = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
          "agosto", "setiembre", "octubre", "noviembre", "diciembre"]


def fecha_larga(valor) -> str:
    """21 de junio de 2009, como en el modelo."""
    if not valor:
        return ""
    try:
        return f"{valor.day:02d} de {_MESES[valor.month - 1]} de {valor.year}"
    except AttributeError:
        return str(valor)[:10]


def edad_de(fecha_nac, hoy=None) -> str:
    if not fecha_nac:
        return ""
    from datetime import date
    hoy = hoy or date.today()
    try:
        nac = fecha_nac.date() if hasattr(fecha_nac, "date") else fecha_nac
        anios = hoy.year - nac.year - ((hoy.month, hoy.day) < (nac.month, nac.day))
        return f"{anios} años"
    except Exception:
        return ""


def preparar_nna(fila: dict) -> dict:
    """Deja el NNA listo para imprimir en la sección I."""
    grado = _GRADO.get(str(fila.get("grado_estudio") or ""), "")
    nivel = _NIVEL.get(str(fila.get("nivel_educativo") or ""), "")
    return {
        **fila,
        "edad": edad_de(fila.get("fecha_nacimiento")),
        "fecha_nacimiento": fecha_larga(fila.get("fecha_nacimiento")),
        "lugar_nacimiento": " - ".join(
            x for x in [fila.get("departamento_nac"), fila.get("distrito_nac")] if x
        ),
        "grado_instruccion": grado or nivel,
    }
